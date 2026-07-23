"""Test _render_table logic độc lập với loader."""
import sys, re
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document as DocxDocument

doc = DocxDocument('data/220269_ Khai pha du lieu - AI.docx')

def _cell_text(cell):
    parts = []
    for para in cell.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for nested in cell.tables:
        for row in nested.rows:
            for c2 in row.cells:
                t2 = c2.text.strip()
                if t2:
                    parts.append(t2)
    return ' '.join(parts)

def _detect_header_row(table):
    if not table.rows: return None
    first = [_cell_text(c).strip() for c in table.rows[0].cells]
    non_empty = [c for c in first if c]
    return first if len(non_empty) >= 2 else None

def _render_table(table):
    if not table.rows: return []
    lines = []
    detected_header = _detect_header_row(table)

    data_start = 0
    if detected_header:
        data_start = 1
        if len(table.rows) > 2:
            row2 = [_cell_text(c).strip() for c in table.rows[1].cells]
            s2 = set(); u2 = []
            for c in row2:
                if c not in s2: u2.append(c); s2.add(c)
            ne2 = set(c for c in u2 if c)
            neh = set(c for c in detected_header if c)
            if ne2 and ne2.issubset(neh):
                data_start = 2

    num_cols = len(table.columns)

    for ri, row in enumerate(table.rows):
        if ri < data_start: continue
        cells = [_cell_text(c).strip() for c in row.cells]
        seen = set(); unique = []
        for c in cells:
            if c not in seen: unique.append(c); seen.add(c)
        non_empty = [c for c in unique if c]
        if not non_empty: continue

        if len(non_empty) == 1:
            lines.append(non_empty[0])
            continue

        if detected_header and len(unique) >= 2:
            parts = []
            for ci, val in enumerate(unique):
                if not val: continue
                hl = detected_header[ci].strip() if ci < len(detected_header) else ''
                if hl and not re.match(r'^[\d.]+$', hl) and hl != val:
                    parts.append(f"{hl}: {val}")
                else:
                    parts.append(val)
            if parts: lines.append(' | '.join(parts))
        elif num_cols == 2 and len(non_empty) == 2:
            lines.append(f"{non_empty[0]}: {non_empty[1]}")
        else:
            lines.append(' | '.join(non_empty))
    return lines


print('=== BEFORE (pipe-join cũ) vs AFTER (semantic) ===\n')

print('--- Table 1: Thông tin học phần ---')
print('[BEFORE] "Loại học phần | Số tín chỉ | Số giờ dự giảng | ..."')
print('[AFTER]')
for line in _render_table(doc.tables[0]):
    print(f'  {line}')

print()
print('--- Table 6: Đánh giá học phần ---')
print('[BEFORE] "Đánh giá quá trình | Kiểm tra lý thuyết | Từ bài 1 đến bài 3 | ..."')
print('[AFTER]')
for line in _render_table(doc.tables[5]):
    print(f'  {line}')

print()
print('--- Table 2: Key-value (Học phần tiên quyết) ---')
print('[BEFORE] "Học phần tiên quyết | Cơ sở dữ liệu"')
print('[AFTER]')
for line in _render_table(doc.tables[1]):
    print(f'  {line}')

print()
print('--- Table 5 (first 6 rows): Nội dung học phần ---')
for line in _render_table(doc.tables[4])[:6]:
    print(f'  {line}')
